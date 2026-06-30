from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt

from src.core.config import settings
from src.models.report import Report


def _output_dir(report_id: str) -> Path:
    out_dir = Path(settings.upload_dir) / "reports" / report_id
    out_dir.mkdir(parents=True, exist_ok=True)
    return out_dir


def build_word(report: Report) -> str:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    soup = BeautifulSoup(report.story_html or "", "html.parser")

    for tag in soup.find_all(["h1", "h2", "p", "ul"]):
        if tag.name == "h1":
            doc.add_heading(tag.get_text(strip=True), level=0)
        elif tag.name == "h2":
            doc.add_heading(tag.get_text(strip=True), level=2)
        elif tag.name == "p":
            text = tag.get_text(strip=True)
            if text:
                doc.add_paragraph(text)
        elif tag.name == "ul":
            for li in tag.find_all("li"):
                doc.add_paragraph(li.get_text(strip=True), style="List Bullet")

    out_path = _output_dir(str(report.id)) / "report.docx"
    doc.save(str(out_path))
    return str(out_path)
